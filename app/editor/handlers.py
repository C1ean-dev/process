import os
import logging
import locale
import re
from datetime import datetime
from flask import render_template, flash, request, current_app, send_file, redirect, url_for
from flask_login import current_user
from docx import Document
from io import BytesIO

from app.models import Group
from .forms import DocumentForm

logger = logging.getLogger(__name__)

# Tenta definir o locale para português para formatar o mês por extenso
try:
    locale.setlocale(locale.LC_TIME, "pt_BR.utf-8")
except:
    try:
        locale.setlocale(locale.LC_TIME, "Portuguese_Brazil.1252")
    except:
        pass

class EditorHandler:
    def list_termos(self):
        termos_dir = os.path.join(current_app.root_path, 'termos')
        termos = []
        if os.path.exists(termos_dir):
            termos = [f for f in os.listdir(termos_dir) if f.endswith('.docx')]
        return render_template('editor/list.html', title='Select Document', termos=termos)

    def edit_termo(self, filename):
        form = DocumentForm()
        user_groups = current_user.groups.all()
        created_groups = Group.query.filter_by(creator_id=current_user.id).all()
        all_groups = list(set(user_groups + created_groups))
        form.group.choices = [(0, 'No Group')] + [(g.id, g.name) for g in all_groups]
        return render_template('editor/edit.html', title='Fill Document Data', filename=filename, form=form)

    def save_termo(self):
        filename = request.form.get('filename')
        form = DocumentForm()
        
        user_groups = current_user.groups.all()
        created_groups = Group.query.filter_by(creator_id=current_user.id).all()
        all_groups = list(set(user_groups + created_groups))
        form.group.choices = [(0, 'No Group')] + [(g.id, g.name) for g in all_groups]

        if not form.validate_on_submit():
            flash('Please fill all required fields.', 'danger')
            return render_template('editor/edit.html', title='Fill Document Data', filename=filename, form=form)

        termos_dir = os.path.join(current_app.root_path, 'termos')
        filepath = os.path.join(termos_dir, filename)

        try:
            doc = Document(filepath)
            
            # Preparar a data atual
            now = datetime.now()
            # Formato: 10 de fevereiro de 2026
            data_extenso = now.strftime("%d de %B de %Y")
            linha_data = f"São Paulo, {data_extenso}."

            # Formatar equipamentos com um tab e hífen antes de cada linha
            equip_data = form.equipamentos.data.strip()
            # Divide por linhas e reconstrói com tab + hífen
            equip_list = [f"\t- {item.strip()}" for item in equip_data.split('\n') if item.strip()]
            equip_formatted = "\n".join(equip_list)

            # Mapeamento usando REGEX super permissivo
            # O padrão r'text.*' tenta casar o inicio do paragrafo.
            replacements = [
                # Match "Empregado:" com qualquer coisa depois
                (re.compile(r'^\s*Empregado:.*', re.IGNORECASE), f"Empregado: {form.nome.data}"),
                
                # Match "Fução:", "Função:", "Fucao:" etc.
                # F, depois u, depois qualquer coisa (como ç, c, n), depois ão/ao, depois :
                (re.compile(r'^\s*Fu.*[çc].*[ãa]o:.*', re.IGNORECASE), f"Função: {form.funcao.data}"),
                
                # Match "Empnregador:", "Empregador:"
                (re.compile(r'^\s*Emp.*[nr].*egador:.*', re.IGNORECASE), f"Empregador: {form.empregador.data}"),
                
                # Match "Matricula:"
                (re.compile(r'^\s*Matr[íi]cula:.*', re.IGNORECASE), f"Matrícula: {form.matricula.data}"),
                
                # Match "R.G. nº:", "RG:", "R.G.:"
                (re.compile(r'^\s*R\.?G\.?.*n?º?:.*', re.IGNORECASE), f"RG: {form.rg.data}"),
                
                # Match "CPF:"
                (re.compile(r'^\s*CPF:.*', re.IGNORECASE), f"CPF: {form.cpf.data}"),
                
                # Match "Descrição dos equipamentos..."
                (re.compile(r'^\s*Descriç.*o dos equipamentos.*', re.IGNORECASE), f"Descrição dos equipamentos/ferramentas:\n{equip_formatted}"),
                
                # Match data: "São Paulo," seguido de qualquer coisa
                (re.compile(r'^\s*São Paulo,.*', re.IGNORECASE), linha_data)
            ]

            self._robust_replace(doc, replacements)

            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            safe_name = "".join([c for c in form.nome.data if c.isalnum() or c in (' ', '_')]).strip().replace(' ', '_')
            download_name = f"Revisao_{safe_name}_{filename}"
            
            return send_file(
                docx_buffer,
                as_attachment=True,
                download_name=download_name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f"Error generating document: {e}", exc_info=True)
            flash(f'An error occurred: {str(e)}', 'danger')
            return redirect(url_for('editor.edit_termo', filename=filename))

    def _robust_replace(self, doc, replacements):
        """Aplica as substituições em parágrafos e tabelas com logs de debug."""
        count = 0
        for p in doc.paragraphs:
            if self._check_and_replace(p, replacements):
                count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if self._check_and_replace(p, replacements):
                            count += 1
        logger.info(f"Total replacements made: {count}")

    def _check_and_replace(self, paragraph, replacements):
        """Verifica e substitui o texto do parágrafo."""
        text = paragraph.text.strip() # Remove espaços extras das pontas para o check
        if not text:
            return False

        for pattern, replacement in replacements:
            # Usamos search para encontrar o padrão no texto original (que pode ter espaços no inicio)
            if pattern.search(paragraph.text):
                logger.info(f"Replacing '{paragraph.text}' with '{replacement}'")
                paragraph.text = replacement
                return True
        return False